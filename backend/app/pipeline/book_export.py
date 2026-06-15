"""
Exporters for the faithful ChatStory manuscript:
  * render_faithful_html  -> A4 HTML (printed to PDF via Playwright/Chromium)
  * build_docx            -> editable .docx with real Word footnotes

Both consume the manuscript dict produced by pipeline/faithful.py. The DOCX is
the editable master (so the reader can fix any inaccuracy); the PDF is the
print-ready A4 version. Footnotes carry the timestamp of each scene so anything
on the page can be cross-referenced against the original export.
"""

from __future__ import annotations

import copy
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape

from ..settings import TEMPLATES_DIR

_jinja = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


# ---------------------------------------------------------------------------
# HTML (for the A4 PDF)
# ---------------------------------------------------------------------------

def render_faithful_html(manuscript: dict, *, preview: bool = False) -> str:
    m = manuscript
    if preview:
        # Free preview = cover + stats + the first chapter only.
        m = copy.deepcopy(manuscript)
        m["chapters"] = m.get("chapters", [])[:1]
    return _jinja.get_template("book_faithful.html").render(m=m, preview=preview)


# ---------------------------------------------------------------------------
# DOCX (editable, with footnotes)
# ---------------------------------------------------------------------------

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_RT_FOOTNOTES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
_CT_FOOTNOTES = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"


class _Footnotes:
    """Attaches a real /word/footnotes.xml part and lets us add Word footnotes.
    If anything goes wrong, `ok` stays False and the caller falls back to inline
    timestamp citations (still valid, still editable)."""

    def __init__(self, document):
        self.ok = False
        self._next_id = 1
        try:
            from docx.oxml import parse_xml
            from docx.opc.part import XmlPart
            from docx.opc.packuri import PackURI

            root = (
                f'<w:footnotes xmlns:w="{_W}">'
                f'<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'</w:footnotes>'
            )
            self._el = parse_xml(root)
            # XmlPart serializes its live element at save time, so footnotes we
            # append after creation are included.
            part = XmlPart(
                PackURI("/word/footnotes.xml"), _CT_FOOTNOTES,
                self._el, document.part.package,
            )
            self._part = part
            document.part.relate_to(part, _RT_FOOTNOTES)
            self.ok = True
        except Exception as exc:  # noqa: BLE001
            print(f"[book_export] real footnotes unavailable, using inline: {exc}")

    def add(self, paragraph, text: str) -> bool:
        """Append a footnote reference to `paragraph` and store its text."""
        if not self.ok:
            return False
        try:
            from docx.oxml import parse_xml
            fid = self._next_id
            self._next_id += 1
            fn = parse_xml(
                f'<w:footnote xmlns:w="{_W}" w:id="{fid}">'
                f'<w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
                f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
                f'<w:r><w:t xml:space="preserve"> {_xml_escape(text)}</w:t></w:r></w:p></w:footnote>'
            )
            self._el.append(fn)
            ref = parse_xml(
                f'<w:r xmlns:w="{_W}"><w:rPr><w:rStyle w:val="FootnoteReference"/>'
                f'<w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="{fid}"/></w:r>'
            )
            paragraph._p.append(ref)
            return True
        except Exception:
            return False


def _xml_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def build_docx(manuscript: dict, out_path: Path) -> bool:
    """Write the manuscript as an editable .docx (A4) with footnotes. Returns
    True on success. Import errors (python-docx missing) return False so the
    pipeline can continue with just the PDF."""
    try:
        from docx import Document
        from docx.shared import Mm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # python-docx not installed
        print(f"[book_export] DOCX skipped (python-docx not available): {exc}")
        return False

    doc = Document()
    # A4 with print-safe margins
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(16)
        section.right_margin = Mm(16)

    footnotes = _Footnotes(doc)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(manuscript.get("title", "Our Story"))
    run.bold = True
    run.font.size = Pt(28)
    sub = doc.add_paragraph(manuscript.get("subtitle", ""))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rng = doc.add_paragraph(manuscript.get("date_range", ""))
    rng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Statistics
    s = manuscript.get("stats", {}) or {}
    doc.add_heading("The Numbers", level=1)
    doc.add_paragraph(f"Messages kept: {s.get('total_messages', 0):,}")
    doc.add_paragraph(f"Days you talked: {s.get('days_active', 0)}")
    doc.add_paragraph(f"From {s.get('first_date','')} to {s.get('last_date','')}")
    if s.get("per_sender"):
        doc.add_paragraph("Who texted more:").runs[0].bold = True
        for who, c in s["per_sender"].items():
            doc.add_paragraph(f"   {who}: {c:,}")
    if s.get("initiations"):
        doc.add_paragraph("Who started conversations more:").runs[0].bold = True
        for who, c in s["initiations"].items():
            doc.add_paragraph(f"   {who}: {c:,}")
    if s.get("per_month"):
        doc.add_paragraph("Messages by month:").runs[0].bold = True
        for row in s["per_month"]:
            doc.add_paragraph(f"   {row['label']}: {row['count']:,}")

    # The book
    grey = RGBColor(0x80, 0x80, 0x80)
    for ch in manuscript.get("chapters", []):
        doc.add_page_break()
        doc.add_heading(ch.get("title", ""), level=1)
        for sc in ch.get("scenes", []):
            setting = sc.get("setting", "")
            p = doc.add_paragraph()
            if setting:
                r = p.add_run(setting)
                r.italic = True
            # Footnote (real if available, else inline citation)
            if not footnotes.add(p, sc.get("footnote", "")):
                cite = p.add_run(f"  [{sc.get('footnote','')}]")
                cite.italic = True
                cite.font.size = Pt(8)
                cite.font.color.rgb = grey
            for l in sc.get("lines", []):
                lp = doc.add_paragraph()
                who = lp.add_run(f"{l['sender']}  ")
                who.bold = True
                lp.add_run(l["text"])
                tm = lp.add_run(f"   {l.get('time','')}")
                tm.font.size = Pt(7.5)
                tm.font.color.rgb = grey
            doc.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return True
